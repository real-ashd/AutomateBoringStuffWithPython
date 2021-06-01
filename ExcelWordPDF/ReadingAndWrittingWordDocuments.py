#Reading and Writting Word Documents
#We use docx module
#To install docx use 'pip install python-docx'
import docx
d = docx.Document('demo.docx')
print(d.paragraphs)			#paragraphs contains list of paragraph objects
print("--------------------------")
#1st paragraph text
print(d.paragraphs[0].text)
print(d.paragraphs[1].text)

#Each paragraph object contains one or more runs
p = d.paragraphs[1]
print("--------------------------")

#This object has 4 run objects
print(p.runs) 			#Like Bold, Italic, normal, underline
print("--------------------------")
print(p.runs[0].text)
print("--------------------------")
print("Bold value is set to :",p.runs[2].bold)
print("Bold value is set to :",p.runs[1].bold)
print("Italic value is set to :",p.runs[2].italic)
print("Italic value is set to :",p.runs[4].italic)
print("--------------------------")

#Changing the values and runs of objects
p.runs[4].underline = True
p.runs[4].text = "italic and underline"

#save it
d.save('demo2.docx')

#Styles of runs
print(p.style)
p.style = 'Title'
print("After changing the style:",p.style)
d.sve('demo3.docx')
print("--------------------------")


#Create a brand new Word Document
d1 = docx.Document()
d1.add_paragraph('Hello this is a paragraph.')
d1.add_paragraph('This is another paragraph.')
p1 = d1.paragraphs[0]
p1.add_run('This is a new run.')	#Second Run
print(p1.runs)
p1.runs[1].bold = True

d1.save('demo4.docx')
print("--------------------------")